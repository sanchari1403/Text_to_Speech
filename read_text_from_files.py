from text_to_speech_converter import text_to_speech
import PyPDF2 
import docx
import doc2text

def read_from_txt(filepath):
    try:
        with open(filepath, 'r') as file:
            txt = file.read()
        if filepath.split(".")[-1]=='txt':
            print(txt)
            text_to_speech(txt)
            file.close()
        else:
            print("The file you uploaded is not a .txt file!")
    except Exception as e:
        print(e)
        print("Kindly upload a proper .txt file.")

def read_from_pdf(filepath):
    try:
        # creating a pdf file object 
        pdfFileObj = open(filepath, 'rb') 
        if filepath.split(".")[-1]=='pdf':    
            # creating a pdf reader object 
            pdfReader = PyPDF2.PdfFileReader(pdfFileObj) 
                
            # getting number of pages in pdf file 
            pages=pdfReader.numPages

            text=''
            # creating a page object
            for i in range(0,pages): 
                pageObj = pdfReader.getPage(i)     
                # extracting text from page 
                text=pageObj.extractText()
                text_to_speech("Reading from page number"+str(i+1)+"of"+str(pages)) 
                text_to_speech(text) 
            pdfFileObj.close() 
        else:
            print("The file you uploaded is not a .pdf file!")
    except Exception as e:
        print(e)
        print("Kindly upload a proper .pdf file.")

def read_from_docx(filepath):
    try:
        doc = docx.Document(filepath)
        if filepath.split(".")[-1]=='docx':  
            fullText = []
            for para in doc.paragraphs:
                fullText.append(para.text)
            text='\n'.join(fullText)
            text_to_speech(text)
        else:
            print("The file you uploaded is not a .docx file!")
    except Exception as e:
        print(e)
        print("Kindly upload a proper .docx file.")

    
def read_from_files(filepath):
    try:
        print("----OCR Processing----")
        doc = doc2text.Document()
        doc = doc2text.Document(lang="eng")
        doc.read(filepath)
        doc.process()
        doc.extract_text()
        text = doc.get_text()
        text_to_speech(text)
    except Exception as e:
        print(e)
        print("Kindly upload a proper file in either of PNG, JPG, JPEG, BMP, TIFF formats.")

    
if __name__ == "__main__":

    try:
        print("******Welcome To Text-To-Speech Converter*****")
        print("Disc: English language supported only")
        print("Disc: Supported file formats are - PDF, TXT, DOCX, PNG, JPG, JPEG, BMP, TIFF")
        choice=input("Type of file: Press 1 for .txt file,\n Press 2 for .pdf file,\n Press 3 for .docx file \n Press 4 for any other file format: ")
        filepath=raw_input("Enter the full filepath: ")
        if choice == 1:
            read_from_txt(filepath)
        elif choice == 2:
            read_from_pdf(filepath)
        elif choice == 3:
            read_from_docx(filepath)
        elif choice == 4:
            read_from_files(filepath)
            print("Thank you for visiting!")
        else:
            print("Please choose from the options!")
    except Exception as e:
        print("Enter choice between 1-4 and full filepath of the file.")